import docx


def add_heading(line, doc):
    i = 0
    while line[i] == '#':
        i += 1
    doc.add_heading(line.strip("# "), i)
    return doc


def markdown_to_docx(text):
    data = text.split("\n")
    file_name = data[0].strip()
    markdown = [x.rstrip() for x in data[1:]]
    doc = docx.Document()
    doc.add_heading(file_name, 0)
    for line in markdown:
        if not line:
            doc.add_paragraph()
        elif line.startswith('#'):
            doc = add_heading(line, doc)
        elif line[0] in '*-+' and line[1] == ' ':
            doc.add_paragraph(line.strip("*-+ "), style='List Bullet')
        elif line[0].isdigit():
            if line.split('.')[0].isdigit():
                doc.add_paragraph(line.rstrip().lstrip("1234567890. "), style='List Number')
            else:
                doc.add_paragraph(line.strip())
        elif line.startswith("*") or line.startswith("_"):
            n = list(line[:3])
            n = n.count(n[0])
            p = doc.add_paragraph()
            run = p.add_run(line.strip("_* "))
            if n in [1, 3]:
                run.italic = True
            if n in [2, 3]:
                run.bold = True
        else:
            doc.add_paragraph(line.strip())
    doc.save('res.docx')
