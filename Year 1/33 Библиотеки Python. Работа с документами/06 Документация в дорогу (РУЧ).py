import docx


def add_heading(line, doc):
    i = 0
    while line[i] == '#':
        i += 1
    doc.add_heading(line.strip("# "), i)
    return doc


def formatting(doc, line, style=''):
    if style:
        p = doc.add_paragraph('', style=style)
    else:
        p = doc.add_paragraph()
    line = line.split()
    bold = False
    italic = False
    for a, i in enumerate(line):
        if i.startswith('___') or i.startswith('***'):
            italic = True
            bold = True
        elif i.startswith('__') or i.startswith('**'):
            bold = True
        elif i.startswith('_') or i.startswith('*'):
            italic = True
        if style == 'List Number':
            run = p.add_run(' ' + i.strip('*_'))
        else:
            run = p.add_run(i.strip('*_') + ' ' if a != len(line) - 1 else i.strip('*_'))
        run.italic = italic
        run.bold = bold
        if i.endswith('___') or i.endswith('***'):
            italic = False
            bold = False
        elif i.endswith('__') or i.endswith('**'):
            bold = False
        elif i.endswith('_') or i.endswith('*'):
            italic = False


def markdown_to_docx(text):
    data = text.split("\n")
    file_name = data[0].strip()
    markdown = [x.rstrip() for x in data[1:]]
    doc = docx.Document()
    doc.add_heading(file_name, 0)
    c = 0
    for line in markdown:
        if not line and c >= 1:
            doc.add_paragraph()
            continue
        elif not line:
            c += 1
            continue
        elif line.startswith('#'):
            doc = add_heading(line, doc)
        elif line[0] in '*-+' and line[1] == ' ':
            formatting(doc, line.strip("*-+"), style='List Bullet')
        elif line[0].isdigit():
            if line.split('.')[0].isdigit():
                formatting(doc, ' ' + line.rstrip().lstrip("1234567890."), style='List Number')
            else:
                formatting(doc, line.strip())
        else:
            formatting(doc, line.strip())
        c = 0
    doc.save(f'{file_name}.docx')
