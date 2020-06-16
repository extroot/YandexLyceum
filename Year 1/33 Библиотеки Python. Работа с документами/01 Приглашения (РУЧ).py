import sys
from docx import Document

place = input()
time = input()
data = list(map(str.strip, sys.stdin))

document = Document()
for x in data:
    document.add_heading("Уважаемая " + x + '!\n')
    p = document.add_paragraph("Приглашаем на мероприятие в честь 8 марта\n")
    p.add_run(place + '\n').italic = True
    p.add_run(time + '\n').bold = True
    document.add_page_break()

document.save('test.docx')
